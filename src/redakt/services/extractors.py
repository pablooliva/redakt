"""Text extraction functions for each supported document format.

Each extractor takes raw file bytes and returns a list of TextChunk dataclass instances.
"""

from __future__ import annotations

import csv
import io
import json
import zipfile
from dataclasses import dataclass, field

import charset_normalizer
import defusedxml
import defusedxml.ElementTree as ET
from bs4 import BeautifulSoup
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.pdfparser import PDFSyntaxError
from striprtf.striprtf import rtf_to_text

from redakt.config import settings


class ExtractionError(Exception):
    """Raised when text extraction fails (corrupt file, password, etc.)."""

    def __init__(self, message: str, status_code: int = 422):
        super().__init__(message)
        self.message = message
        self.status_code = status_code


@dataclass
class TextChunk:
    text: str
    chunk_id: str  # e.g., "A1", "page_1", "paragraph_3"
    chunk_type: str  # e.g., "cell", "page", "paragraph"


@dataclass
class ExtractionResult:
    chunks: list[TextChunk]
    metadata: dict = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _decode_text(raw: bytes) -> str:
    """Decode bytes to string: try UTF-8, then charset-normalizer."""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass

    result = charset_normalizer.from_bytes(raw).best()
    if result is None or result.encoding is None:
        raise ExtractionError(
            "Could not determine file encoding. Please save the file as UTF-8 and re-upload.",
            status_code=422,
        )
    # charset_normalizer doesn't expose a simple confidence float on the result object;
    # we check coherence instead (0.0 = incoherent, 1.0 = perfectly coherent).
    # A low coherence suggests unreliable detection.
    if result.coherence < 0.5:
        raise ExtractionError(
            "Could not determine file encoding. Please save the file as UTF-8 and re-upload.",
            status_code=422,
        )
    return str(result)


def _check_zip_bomb(raw: bytes) -> None:
    """Reject ZIP archives whose total uncompressed size exceeds the limit."""
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            total = sum(info.file_size for info in zf.infolist())
            if total > settings.max_zip_uncompressed_size:
                raise ExtractionError(
                    "The file's compressed content is too large to process safely. "
                    "Please use a simpler document.",
                    status_code=400,
                )
    except zipfile.BadZipFile:
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )


# ---------------------------------------------------------------------------
# Plain text formats
# ---------------------------------------------------------------------------

def extract_txt(raw: bytes) -> ExtractionResult:
    """Extract plain text (.txt) files."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})
    text = _decode_text(raw)
    chunks = [TextChunk(text=text, chunk_id="content", chunk_type="text")]
    return ExtractionResult(chunks=chunks, metadata={})


def extract_md(raw: bytes) -> ExtractionResult:
    """Extract Markdown (.md) files. Treated as plain text (formatting preserved)."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})
    text = _decode_text(raw)
    chunks = [TextChunk(text=text, chunk_id="content", chunk_type="text")]
    return ExtractionResult(chunks=chunks, metadata={})


# ---------------------------------------------------------------------------
# Structured text formats
# ---------------------------------------------------------------------------

def extract_csv(raw: bytes) -> ExtractionResult:
    """Extract CSV files cell-by-cell with auto-detected delimiter."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={"delimiter": ","})
    text = _decode_text(raw)

    # Auto-detect delimiter
    try:
        dialect = csv.Sniffer().sniff(text[:8192])
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ","

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    chunks: list[TextChunk] = []
    cell_count = 0
    for row_idx, row in enumerate(reader, start=1):
        for col_idx, cell in enumerate(row):
            cell_count += 1
            chunk_id = f"R{row_idx}C{col_idx + 1}"
            chunks.append(TextChunk(text=cell, chunk_id=chunk_id, chunk_type="cell"))

    return ExtractionResult(
        chunks=chunks,
        metadata={"delimiter": delimiter, "cells_processed": cell_count},
    )


def extract_json(raw: bytes) -> ExtractionResult:
    """Extract string values from JSON files recursively."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})
    text = _decode_text(raw)

    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ExtractionError(
            f"The file could not be parsed. It may be corrupted or in an unsupported variant. ({exc})",
            status_code=422,
        )

    chunks: list[TextChunk] = []
    _extract_json_strings(data, "", chunks, depth=0)
    return ExtractionResult(chunks=chunks, metadata={"original_structure": data})


_MAX_JSON_DEPTH = 100


def _extract_json_strings(
    obj: object, path: str, chunks: list[TextChunk], depth: int = 0
) -> None:
    """Recursively extract string values from parsed JSON."""
    if depth > _MAX_JSON_DEPTH:
        raise ExtractionError(
            f"JSON nesting exceeds the maximum depth of {_MAX_JSON_DEPTH}. "
            "Please simplify the file structure.",
            status_code=422,
        )
    if isinstance(obj, str):
        chunks.append(TextChunk(text=obj, chunk_id=path or "$", chunk_type="value"))
    elif isinstance(obj, dict):
        for key, value in obj.items():
            _extract_json_strings(value, f"{path}.{key}" if path else f"$.{key}", chunks, depth + 1)
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            _extract_json_strings(item, f"{path}[{i}]", chunks, depth + 1)
    # Numbers, booleans, None — skip (preserved in output)


# ---------------------------------------------------------------------------
# XML / HTML
# ---------------------------------------------------------------------------

def extract_xml(raw: bytes) -> ExtractionResult:
    """Extract text nodes from XML files using defusedxml."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})

    try:
        root = ET.fromstring(raw)
    except (ET.ParseError, defusedxml.DefusedXmlException) as exc:
        raise ExtractionError(
            f"The file could not be parsed. It may be corrupted or in an unsupported variant. ({exc})",
            status_code=422,
        )

    texts: list[str] = []
    _collect_xml_text(root, texts)
    combined = " ".join(texts)
    if not combined.strip():
        return ExtractionResult(chunks=[], metadata={})

    chunks = [TextChunk(text=combined, chunk_id="content", chunk_type="text")]
    return ExtractionResult(chunks=chunks, metadata={})


def _collect_xml_text(element, texts: list[str]) -> None:
    """Walk an XML tree and collect all text content."""
    if element.text and element.text.strip():
        texts.append(element.text.strip())
    for child in element:
        _collect_xml_text(child, texts)
        if child.tail and child.tail.strip():
            texts.append(child.tail.strip())


def extract_html(raw: bytes) -> ExtractionResult:
    """Extract text from HTML files using BeautifulSoup, stripping script/style."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})

    soup = BeautifulSoup(raw, "html.parser")
    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator=" ", strip=True)
    if not text.strip():
        return ExtractionResult(chunks=[], metadata={})

    chunks = [TextChunk(text=text, chunk_id="content", chunk_type="text")]
    return ExtractionResult(chunks=chunks, metadata={})


# ---------------------------------------------------------------------------
# Binary formats (require library imports)
# ---------------------------------------------------------------------------

def extract_xlsx(raw: bytes) -> ExtractionResult:
    """Extract string cells from XLSX workbooks."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={"sheets_processed": 0, "cells_processed": 0})

    _check_zip_bomb(raw)

    import openpyxl  # imported here so defusedxml.defuse_stdlib() runs first

    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypted" in msg:
            raise ExtractionError(
                "This file appears to be password-protected. "
                "Please remove the password and re-upload.",
                status_code=422,
            )
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )

    chunks: list[TextChunk] = []
    sheets_data: dict[str, list] = {}
    text_cell_count = 0

    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows: list[list] = []
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0

            for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
                row_data: list = []
                for cell in row:
                    value = cell.value
                    if isinstance(value, str) and value.strip():
                        text_cell_count += 1
                        if text_cell_count > settings.max_xlsx_cells:
                            wb.close()
                            raise ExtractionError(
                                f"This spreadsheet contains too many cells to process "
                                f"(limit: {settings.max_xlsx_cells:,} text cells). "
                                f"Please reduce the data or split into smaller files.",
                                status_code=422,
                            )
                        # Build cell reference
                        col_letter = _col_num_to_letter(cell.column)
                        chunk_id = f"{sheet_name}!{col_letter}{cell.row}"
                        chunks.append(
                            TextChunk(text=value, chunk_id=chunk_id, chunk_type="cell")
                        )
                        row_data.append(value)
                    else:
                        # Preserve non-string values as-is for output
                        row_data.append(value)
                sheet_rows.append(row_data)
            sheets_data[sheet_name] = sheet_rows
    finally:
        wb.close()

    return ExtractionResult(
        chunks=chunks,
        metadata={
            "sheets_processed": len(sheets_data),
            "cells_processed": text_cell_count,
            "sheets_data": sheets_data,
        },
    )


def _col_num_to_letter(col: int) -> str:
    """Convert 1-based column number to Excel letter (1->A, 27->AA)."""
    result = ""
    while col > 0:
        col, remainder = divmod(col - 1, 26)
        result = chr(65 + remainder) + result
    return result


def extract_docx(raw: bytes) -> ExtractionResult:
    """Extract text from DOCX files (paragraphs + table cells)."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})

    _check_zip_bomb(raw)

    import docx  # imported here so defusedxml.defuse_stdlib() runs first

    try:
        doc = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypted" in msg:
            raise ExtractionError(
                "This file appears to be password-protected. "
                "Please remove the password and re-upload.",
                status_code=422,
            )
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )

    chunks: list[TextChunk] = []
    para_idx = 0

    for para in doc.paragraphs:
        para_idx += 1
        if para.text.strip():
            chunks.append(
                TextChunk(
                    text=para.text,
                    chunk_id=f"paragraph_{para_idx}",
                    chunk_type="paragraph",
                )
            )

    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        for row_idx, row in enumerate(table.rows, start=1):
            for col_idx, cell in enumerate(row.cells, start=1):
                if cell.text.strip():
                    chunks.append(
                        TextChunk(
                            text=cell.text,
                            chunk_id=f"table_{table_idx}_R{row_idx}C{col_idx}",
                            chunk_type="table_cell",
                        )
                    )

    return ExtractionResult(chunks=chunks, metadata={})


def extract_rtf(raw: bytes) -> ExtractionResult:
    """Extract plain text from RTF files using striprtf."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={})

    try:
        text_content = raw.decode("utf-8", errors="replace")
        text = rtf_to_text(text_content)
    except Exception:
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )

    if not text.strip():
        return ExtractionResult(chunks=[], metadata={})

    chunks = [TextChunk(text=text, chunk_id="content", chunk_type="text")]
    return ExtractionResult(chunks=chunks, metadata={})


def extract_pdf(raw: bytes) -> ExtractionResult:
    """Extract text from PDF files using pdfminer.six."""
    if not raw:
        return ExtractionResult(chunks=[], metadata={"pages_processed": 0})

    try:
        text = pdfminer_extract_text(io.BytesIO(raw))
    except PDFSyntaxError:
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )
    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypted" in msg:
            raise ExtractionError(
                "This file appears to be password-protected. "
                "Please remove the password and re-upload.",
                status_code=422,
            )
        raise ExtractionError(
            "The file could not be parsed. It may be corrupted or in an unsupported variant.",
            status_code=422,
        )

    warnings: list[str] = []
    # Heuristic: limited text from a large file suggests scanned/image PDF
    if len(text.strip()) < 100 and len(raw) > 10_000:
        warnings.append(
            "Limited text could be extracted from this PDF. Results may be incomplete. "
            "Scanned or image-based PDFs are not supported in this version."
        )

    if not text.strip():
        return ExtractionResult(chunks=[], metadata={"pages_processed": 0}, warnings=warnings)

    # Split by page (pdfminer uses form-feed characters between pages)
    pages = text.split("\x0c")
    chunks: list[TextChunk] = []
    for i, page_text in enumerate(pages, start=1):
        if page_text.strip():
            chunks.append(
                TextChunk(text=page_text, chunk_id=f"page_{i}", chunk_type="page")
            )

    return ExtractionResult(
        chunks=chunks,
        metadata={"pages_processed": len(chunks)},
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Extractor registry
# ---------------------------------------------------------------------------

EXTRACTORS: dict[str, callable] = {
    ".txt": extract_txt,
    ".md": extract_md,
    ".csv": extract_csv,
    ".json": extract_json,
    ".xml": extract_xml,
    ".html": extract_html,
    ".xlsx": extract_xlsx,
    ".docx": extract_docx,
    ".rtf": extract_rtf,
    ".pdf": extract_pdf,
}
